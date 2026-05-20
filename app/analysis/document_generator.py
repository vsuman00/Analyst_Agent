"""
document_generator.py — Layer 3 Tool
--------------------------------------
DocumentGenerator (.docx)

Converts a final BRD Markdown file to a professionally formatted .docx document.

Input:
  - brd_markdown : str  (raw Markdown text of the final BRD)

Output:
  - brd_output.docx  (written to --out path or current working directory)

Markdown → .docx Mapping:
  # Title         → Title style (bold, 22pt)
  ## Heading      → Heading 1
  ### Subheading  → Heading 2
  #### Subsub     → Heading 3
  | Table |       → docx Table (styled, with header row shading)
  - Bullet        → List Bullet style
  1. Numbered     → List Number style
  **bold**        → inline bold run
  *italic* / _it_ → inline italic run
  `code`          → inline monospace run (Courier New)
  ---             → paragraph page-separator rule (no new page)
  Plain text      → Normal paragraph

Dependencies:
  python-docx>=1.1.0  (install: pip install python-docx)

Usage:
  python3 -m app.analysis.document_generator \\
    --brd runtime/outputs/brd.md \\
    --out runtime/outputs/brd_output.docx
"""

from __future__ import annotations

import re
import sys
import argparse
from pathlib import Path
from typing import List, Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print(
        "[ERROR] python-docx is not installed.\n"
        "        Install it with: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)

# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

FONT_BODY       = "Calibri"
FONT_HEADING    = "Calibri"
FONT_MONO       = "Courier New"

SIZE_TITLE      = 22
SIZE_HEADING1   = 16
SIZE_HEADING2   = 13
SIZE_BODY       = 11

COLOR_HEADING1  = RGBColor(0x1F, 0x49, 0x7D)   # Deep navy
COLOR_HEADING2  = RGBColor(0x2E, 0x74, 0xB5)   # Mid blue
COLOR_TABLE_HDR = RGBColor(0x1F, 0x49, 0x7D)   # Navy (table header bg)
COLOR_DIVIDER   = RGBColor(0xCC, 0xCC, 0xCC)   # Light grey

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell, fill_hex: str):
    """Apply a solid background shading to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _add_hr(doc: Document):
    """Add a light-grey horizontal rule paragraph."""
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)


def _apply_inline_markup(run_container, text: str):
    """
    Parse inline markdown and add styled runs:
      **bold**, *italic*, _italic_, `code`, plain text.
    """
    # Order matters: bold (**...**) must be tried before single-asterisk italic.
    # Plain-text group must NOT consume lone '*' so the italic branch can claim it.
    token_re = re.compile(
        r'\*\*(.+?)\*\*'              # group 1: **bold**
        r'|\*([^*\n]+?)\*'            # group 2: *italic*
        r'|_([^_\n]+?)_'              # group 3: _italic_
        r'|`([^`\n]+?)`'              # group 4: `code`
        r'|([^*_`]+)'                 # group 5: plain text (no *, _, `)
    )
    for m in token_re.finditer(text):
        bold, ast_it, und_it, code, plain = m.groups()
        if bold is not None:
            r = run_container.add_run(bold)
            r.bold = True
            r.font.name = FONT_BODY
            r.font.size = Pt(SIZE_BODY)
        elif ast_it is not None or und_it is not None:
            r = run_container.add_run(ast_it if ast_it is not None else und_it)
            r.italic = True
            r.font.name = FONT_BODY
            r.font.size = Pt(SIZE_BODY)
        elif code is not None:
            r = run_container.add_run(code)
            r.font.name = FONT_MONO
            r.font.size = Pt(SIZE_BODY - 1)
        elif plain is not None:
            r = run_container.add_run(plain)
            r.font.name = FONT_BODY
            r.font.size = Pt(SIZE_BODY)


def _parse_table_row(line: str) -> List[str]:
    """Parse a Markdown table row into a list of cell strings."""
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_separator_row(line: str) -> bool:
    """Return True if the line is a Markdown table separator (|---|---|)."""
    return bool(re.match(r'^\|[\s\-|:]+\|$', line.strip()))


# ---------------------------------------------------------------------------
# Core renderer
# ---------------------------------------------------------------------------

def markdown_to_docx(brd_markdown: str, out_path: Path) -> None:
    doc = Document()

    # Set narrow-ish margins for a professional look
    section = doc.sections[0]
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    lines = brd_markdown.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Table detection ──────────────────────────────────────────────
        # A table starts with a '|' line followed by a separator row
        if (
            line.strip().startswith("|")
            and i + 1 < len(lines)
            and _is_separator_row(lines[i + 1])
        ):
            header_cells = _parse_table_row(line)
            col_count = len(header_cells)

            # Collect all body rows (skip separator)
            body_rows: List[List[str]] = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                if not _is_separator_row(lines[j]):
                    body_rows.append(_parse_table_row(lines[j]))
                j += 1

            row_count = 1 + len(body_rows)
            tbl = doc.add_table(rows=row_count, cols=col_count)
            tbl.style = "Table Grid"

            # Header row
            hdr_row = tbl.rows[0]
            for ci, cell_text in enumerate(header_cells):
                cell = hdr_row.cells[ci]
                _shade_cell(cell, "1F497D")
                p = cell.paragraphs[0]
                r = p.add_run(cell_text.replace("**", "").replace("_", ""))
                r.bold       = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.name  = FONT_HEADING
                r.font.size  = Pt(SIZE_BODY)

            # Body rows
            for ri, row_cells in enumerate(body_rows):
                tbl_row = tbl.rows[ri + 1]
                for ci, cell_text in enumerate(row_cells[:col_count]):
                    cell = tbl_row.cells[ci]
                    p    = cell.paragraphs[0]
                    _apply_inline_markup(p, cell_text)
                    # Alternate row shading for readability
                    if ri % 2 == 1:
                        _shade_cell(cell, "EEF3F9")

            doc.add_paragraph()   # breathing room after table
            i = j
            continue

        # ── Horizontal rule ──────────────────────────────────────────────
        if line.strip() in ("---", "***", "___"):
            _add_hr(doc)
            i += 1
            continue

        # ── H1 — Document title ──────────────────────────────────────────
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(text)
            r.bold            = True
            r.font.name       = FONT_HEADING
            r.font.size       = Pt(SIZE_TITLE)
            r.font.color.rgb  = COLOR_HEADING1
            i += 1
            continue

        # ── H2 — Section headings ────────────────────────────────────────
        if line.startswith("## ") and not line.startswith("### "):
            text = line[3:].strip()
            p = doc.add_heading(level=1)
            p.clear()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            r.bold            = True
            r.font.name       = FONT_HEADING
            r.font.size       = Pt(SIZE_HEADING1)
            r.font.color.rgb  = COLOR_HEADING1
            i += 1
            continue

        # ── H3 — Subsection headings ─────────────────────────────────────
        if line.startswith("### ") and not line.startswith("#### "):
            text = line[4:].strip()
            p = doc.add_heading(level=2)
            p.clear()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(text)
            r.bold            = True
            r.font.name       = FONT_HEADING
            r.font.size       = Pt(SIZE_HEADING2)
            r.font.color.rgb  = COLOR_HEADING2
            i += 1
            continue

        # ── H4 — Sub-subsection headings ─────────────────────────────────
        if line.startswith("#### "):
            text = line[5:].strip()
            p = doc.add_heading(level=3)
            p.clear()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(text)
            r.bold            = True
            r.font.name       = FONT_HEADING
            r.font.size       = Pt(SIZE_BODY + 1)
            r.font.color.rgb  = RGBColor(0x40, 0x40, 0x40)
            i += 1
            continue

        # ── Bullet points ────────────────────────────────────────────────
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent   = Inches(0.25)
            p.paragraph_format.space_after   = Pt(2)
            _apply_inline_markup(p, text)
            i += 1
            continue

        # ── Numbered list items (1. Item) ────────────────────────────────
        numbered = re.match(r'^(\d+)\.\s+(.*)', line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(2)
            _apply_inline_markup(p, numbered.group(2).strip())
            i += 1
            continue

        # ── Metadata key-value lines (bold label: value) ─────────────────
        # e.g.  "**Document Type:** Business Requirement Document  "
        bold_kv = re.match(r'^\*\*(.+?)\*\*\s*[:\u2014\-]\s*(.*)', line)
        if bold_kv:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            label_run = p.add_run(bold_kv.group(1) + ": ")
            label_run.bold       = True
            label_run.font.name  = FONT_BODY
            label_run.font.size  = Pt(SIZE_BODY)
            _apply_inline_markup(p, bold_kv.group(2).rstrip())
            i += 1
            continue

        # ── Blank / whitespace-only line ─────────────────────────────────
        if not line.strip():
            # Only add paragraph gap if last paragraph was non-empty
            i += 1
            continue

        # ── Plain / mixed inline markup paragraph ────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        _apply_inline_markup(p, line.strip())
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[OK] Document saved to {out_path}", file=sys.stderr)


# ---------------------------------------------------------------------------
# CLI entrypoint
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="DocumentGenerator: Convert BRD Markdown to professional .docx"
    )
    parser.add_argument(
        "--brd",
        required=True,
        help="Path to the final BRD Markdown file",
    )
    parser.add_argument(
        "--out",
        default="runtime/outputs/brd_output.docx",
        help="Output .docx path (default: runtime/outputs/brd_output.docx)",
    )
    args = parser.parse_args()

    brd_path = Path(args.brd)
    if not brd_path.exists():
        print(f"[ERROR] BRD file not found: {brd_path}", file=sys.stderr)
        sys.exit(1)

    brd_text = brd_path.read_text(encoding="utf-8")
    markdown_to_docx(brd_text, Path(args.out))
